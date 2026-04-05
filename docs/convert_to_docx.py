"""
Convert mid-term-report.md to mid-term-report.docx
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR = Path(__file__).parent
MD_FILE  = DOCS_DIR / "mid-term-report.md"
DOCX_FILE = DOCS_DIR / "mid-term-report.docx"
BASE_DIR = DOCS_DIR.parent   # F:/Thesis

# ── helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def resolve_image(path_str: str) -> Path | None:
    """Resolve image path relative to the docs directory."""
    p = (DOCS_DIR / path_str).resolve()
    if p.exists():
        return p
    # try relative to base
    p2 = (BASE_DIR / path_str.lstrip("../")).resolve()
    if p2.exists():
        return p2
    return None

def strip_bold(text: str) -> list[tuple[str, bool]]:
    """Return list of (text, is_bold) tuples."""
    parts = []
    for seg in re.split(r'(\*\*[^*]+\*\*)', text):
        if seg.startswith("**") and seg.endswith("**"):
            parts.append((seg[2:-2], True))
        elif seg:
            parts.append((seg, False))
    return parts

def add_run_styled(para, raw_text: str, bold=False, italic=False, code=False, font_size=10):
    """Add a run with inline **bold** / `code` parsing."""
    segments = strip_bold(raw_text)
    for seg_text, seg_bold in segments:
        # Handle inline code
        code_parts = re.split(r'(`[^`]+`)', seg_text)
        for cp in code_parts:
            if cp.startswith("`") and cp.endswith("`"):
                r = para.add_run(cp[1:-1])
                r.font.name = "Courier New"
                r.font.size = Pt(font_size - 1)
                r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            elif cp:
                r = para.add_run(cp)
                r.bold = bold or seg_bold
                r.italic = italic
                r.font.size = Pt(font_size)


# ── style setup ─────────────────────────────────────────────────────────────

def apply_styles(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    for lvl, name in [(1, "Heading 1"), (2, "Heading 2"),
                      (3, "Heading 3"), (4, "Heading 4")]:
        s = doc.styles[name]
        s.font.name = "Times New Roman"
        s.font.bold = True
        sizes = {1: 16, 2: 14, 3: 12, 4: 11}
        s.font.size = Pt(sizes[lvl])
        colors = {1: RGBColor(0x1a, 0x1a, 0x2e),
                  2: RGBColor(0x16, 0x21, 0x3e),
                  3: RGBColor(0x0f, 0x3c, 0x78),
                  4: RGBColor(0x2c, 0x3e, 0x50)}
        s.font.color.rgb = colors[lvl]
        s.paragraph_format.space_before = Pt(sizes[lvl] - 4)
        s.paragraph_format.space_after  = Pt(4)

    # Code block style
    if "Code Block" not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style("Code Block", 1)
        code_style.base_style = doc.styles["Normal"]
        code_style.font.name = "Courier New"
        code_style.font.size = Pt(9)
        code_style.paragraph_format.left_indent = Cm(1)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after  = Pt(4)

    # Caption style
    if "Fig Caption" not in [s.name for s in doc.styles]:
        cap = doc.styles.add_style("Fig Caption", 1)
        cap.base_style = doc.styles["Normal"]
        cap.font.italic = True
        cap.font.size = Pt(9)
        cap.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)


# ── table builder ────────────────────────────────────────────────────────────

def add_table(doc, rows):
    """rows = list of list of strings; first row = header."""
    num_cols = len(rows[0])
    tbl = doc.add_table(rows=0, cols=num_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = tbl.add_row()
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            is_header = (i == 0)
            if is_header:
                set_cell_bg(cell, "1a1a2e")
                r = p.add_run(cell_text.strip().strip("*"))
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9)
            else:
                # Alternate row shading
                bg = "f0f4f8" if i % 2 == 0 else "ffffff"
                set_cell_bg(cell, bg)
                raw = cell_text.strip()
                add_run_styled(p, raw, font_size=9)

    doc.add_paragraph()  # spacing after table
    return tbl


# ── title page ───────────────────────────────────────────────────────────────

def add_title_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Mid-Term Check Report")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle.add_run("Research on Industrial Defect Detection Methods\nBased on Deep Learning")
    r2.bold = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0x0f, 0x3c, 0x78)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = [
        ("Student Name", "Mohammad Hamim"),
        ("Major", "Computer Science and Technology"),
        ("Supervisor", "Lu Yang (卢洋)"),
        ("Department", "School of Computer Science and Artificial Intelligence"),
        ("University", "Zhengzhou University"),
        ("Report Date", "April 2026"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_lbl = p.add_run(f"{label}: ")
        r_lbl.bold = True
        r_lbl.font.size = Pt(11)
        r_val = p.add_run(value)
        r_val.font.size = Pt(11)

    doc.add_page_break()


# ── markdown parser ───────────────────────────────────────────────────────────

def process_md(doc, md_text: str):
    """Parse markdown lines and add to doc."""
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lines = []

    skip_title = True  # skip the first h1 (used as title page)

    while i < len(lines):
        line = lines[i]

        # ── code block ─────────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                for cl in code_lines:
                    p = doc.add_paragraph(style="Code Block")
                    p.add_run(cl)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── horizontal rule ────────────────────────────────────
        if re.match(r'^-{3,}$', line.strip()):
            i += 1
            continue

        # ── headings ───────────────────────────────────────────
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text  = m.group(2).strip()
            # Skip document title (already on title page)
            if level == 1 and skip_title and "Mid-Term" in text:
                skip_title = False
                i += 1
                continue
            if level == 1 and "Research on Industrial" in text:
                i += 1
                continue
            style_name = f"Heading {min(level, 4)}"
            p = doc.add_paragraph(style=style_name)
            # Strip markdown bold from heading
            clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            p.add_run(clean)
            i += 1
            continue

        # ── image ──────────────────────────────────────────────
        m_img = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if m_img:
            alt   = m_img.group(1)
            src   = m_img.group(2)
            img_p = resolve_image(src)
            if img_p:
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(img_p), width=Inches(5.5))
                except Exception as e:
                    p = doc.add_paragraph(f"[Figure: {alt}]")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_paragraph(f"[Figure: {alt} — image not found at {src}]")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].italic = True
            i += 1
            continue

        # ── italic caption line  (*text*) ──────────────────────
        m_cap = re.match(r'^\*([^*]+)\*\s*$', line.strip())
        if m_cap:
            p = doc.add_paragraph(style="Fig Caption")
            p.add_run(m_cap.group(1))
            i += 1
            continue

        # ── table ──────────────────────────────────────────────
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                raw = lines[i].strip()
                # skip separator rows
                if not re.match(r'^\|[-| :]+\|$', raw):
                    cells = [c.strip() for c in raw.strip("|").split("|")]
                    table_lines.append(cells)
                i += 1
            if table_lines:
                add_table(doc, table_lines)
            continue

        # ── bullet / numbered list ─────────────────────────────
        m_bullet = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.*)', line)
        if m_bullet:
            indent = len(m_bullet.group(1)) // 2
            text   = m_bullet.group(3)
            p = doc.add_paragraph(style="List Bullet" if indent == 0 else "List Bullet 2")
            p.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
            add_run_styled(p, text, font_size=10)
            i += 1
            continue

        # ── blank line ─────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── regular paragraph ──────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_run_styled(p, line.strip(), font_size=11)
        i += 1


# ── main ─────────────────────────────────────────────────────────────────────

def main():
    md_text = MD_FILE.read_text(encoding="utf-8")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    apply_styles(doc)
    add_title_page(doc)
    process_md(doc, md_text)

    doc.save(str(DOCX_FILE))
    print(f"✅  Saved: {DOCX_FILE}")

if __name__ == "__main__":
    main()
